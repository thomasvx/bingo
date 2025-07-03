from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
import random
from random import sample
from math import floor
import statistics
import os
from tkinter import filedialog
from difflib import SequenceMatcher
from bingo_itembanks import *

#Settings (e.g. difficulty, time to bingo)
ROWS = 5
COLUMNS = 5
CARDS_NUMBER = 120
NUMBER_OF_ITEMS_TO_CROP_ITEM_BANK_TO = 0            # 0 to not crop
isMiddleFree = True
TOTAL_NUMBER_OF_FREE_SPACES = 0
ANALYSIS_PERCENT_UNTIL_BINGO = 20
ANALYSIS_BINGO_MODE = "line bingo"                  # "line bingo" or "full card bingo"

#Forced items (if fewer than NUMBER_OF_ITEMS_TO_CROP_ITEM_BANK_TO, otherwise prioritised)
forced = []

item_bank = []
selections_to_include_in_item_bank = [
    thomas_ultraEZBingo250701_artistsSong
]

#Layout
COLOR_FONT = "363636"
COLOR_CELL_BACKGROUND = "F2F2F2"
COLOR_CELL_BORDER = "A1A1A1"

selections_to_prioritise_in_item_bank = [forced, ]

for selection in selections_to_include_in_item_bank:
    for item in selection:
        item_bank.append(item)

def remove_full_duplicates(item_bank):
    l = []
    for item in item_bank:
        if item_bank.count(item) > 1:
            l.append(item)

    if l:
        l = list(set(l))
        l.sort()
        print(f"{len(l)} duplicate(s) removed: ", end="")    
        if len(l) == 1:
            print(f"'{l[0]}'.")
        else:
            for item in l:
                if item == l[-1]:
                    print(f"and '{item}'.")
                else:
                    print(f"'{item}'", end=", ")
            
        return list(set(item_bank))
    else:
        return item_bank

def report_partial_duplicates(item_bank):
    checked_item_bank = []
    for item in item_bank:
        for existing_item in checked_item_bank:
            res = SequenceMatcher(None, item, existing_item).ratio()
            if item in existing_item or existing_item in item:
                print(f"❓ The item '{item}' is pretty similar to '{existing_item}'.")
            elif res > .8:
                print(f"❓ The item '{item}' is pretty similar to '{existing_item}'.")
        checked_item_bank.append(item)
    return

def generate_bingo_card(item_bank):
    bingo_card = []
    placeholder_items = []

    # Fill card with placeholder items
    for _ in range(ROWS * COLUMNS):
        placeholder_items.append("")

    for i in range(COLUMNS):
        lower_bound = ROWS * i
        upper_bound = ROWS * (i + 1)
        row = placeholder_items[lower_bound:upper_bound]

        bingo_card.append(row)
    
    #Add free spaces
    if isMiddleFree:
        if ROWS % 2 == 1 and COLUMNS % 2 == 1:
            middle_row = int((ROWS - 1) / 2)
            middle_column = int((COLUMNS - 1) / 2)
            bingo_card[middle_column][middle_row] = "🆓"

    for _ in range(TOTAL_NUMBER_OF_FREE_SPACES):
        placed = False
        while not placed:
            r = random.randint(0, ROWS - 1)
            s = random.randint(0, COLUMNS - 1)
            if bingo_card[r][s] != "🆓":
                bingo_card[r][s] = "🆓"
                placed = True

    # Fill placeholder items
    if TOTAL_NUMBER_OF_FREE_SPACES == 0 and isMiddleFree:
        frees = 1
    else:
        frees = TOTAL_NUMBER_OF_FREE_SPACES

    items = sample(item_bank, ROWS * COLUMNS - frees)
    random.shuffle(items)

    for row in bingo_card:
        i = bingo_card.index(row)
        for cell in row:
            j = row.index(cell)
            if cell == "":
                bingo_card[i][j] = items.pop()

    return list(map(list, zip(*bingo_card)))

def set_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4.5') #half px
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), COLOR_CELL_BORDER) #gray
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, **kwargs):
    """
    cell:  actual cell instance you want to modify
    usage:
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)

def set_cell_background(cell, fill, color=None, val=None):
    """
    @fill: Specifies the color to be used for the background
    @color: Specifies the color to be used for any foreground
    pattern specified with the val attribute
    @val: Specifies the pattern to be used to lay the pattern
    color over the background color.

    # pip install python-docx==0.8.11

    """
    from docx.oxml.shared import qn  # feel free to move these out
    from docx.oxml.xmlchemy import OxmlElement

    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd') # add new w:shd element to it
    if fill:
        cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
    if color:
        pass # TODO
    if val:
        pass # TODO
    cell_properties.append(cell_shading)  # finally extend cell props with shading element

#Analysis
def isLineBingo(card):
    # row bingo
    for row in card:
        free = 0
        for item in row:
            if item == "🆓":
                free += 1
        if free == ROWS:
            return True
    
    # column bingo
    for column in range(COLUMNS):
        free = 0
        for row in card:
            if row[column] == "🆓":
                free += 1
        if free == COLUMNS:
            return True
    
    # diagonal bingo
    if card[0][0] == "🆓":#TODO custom size
        if card[1][1] == "🆓":
            if card[2][2] == "🆓":
                if card[3][3] == "🆓":
                    if card[4][4] == "🆓":
                        return True
    if card[0][4] == "🆓":
        if card[1][3] == "🆓":
            if card[2][2] == "🆓":
                if card[3][1] == "🆓":
                    if card[4][0] == "🆓":
                        return True
    return False

def isFullBingo(card):
    free = 0
    for row in card:
        for item in row:
            if item == "🆓":
                free += 1
    if free == (ROWS * COLUMNS):
        return True
    else:
        return False

def delCalloutFromCards(callout, all_cards):
    for card in all_cards:
        for row in card:
            for word in row:
                if word == callout:
                    index = row.index(callout)
                    row[index] = "🆓"
    return

def simulateUntilBingo(cards, all_items_on_cards, mode="line bingo"):
    copied_cards = cards.copy()
    l = all_items_on_cards.copy()

    bingo_callouts = []
    bingo = False
    while not bingo:
        r = random.choice(l)
        l.remove(r)
        bingo_callouts.append(r)

        delCalloutFromCards(r, copied_cards)
        m=[]
        for card in copied_cards:
            callout_amount = len(all_items_on_cards) - len(l)
            if mode == "line bingo" and isLineBingo(card):
                bingo = True
                # if (callout_amount == 4 and isMiddleFree) or (callout_amount == 5 and not isMiddleFree):
                #     print(f'I simulated a perfect callout sequence for a {mode} after just {callout_amount} callouts.')
            if mode == "full card bingo" and isFullBingo(card):
                bingo = True
                # if (callout_amount == 24 and isMiddleFree) or (callout_amount == 25 and not isMiddleFree):
                    # print(f'I simulated a perfect callout sequence for a {mode} after just {callout_amount} callouts.')
    return len(bingo_callouts)

def simulateUntilBingoPercent(cards, all_items_on_cards, percent, mode="line bingo"):
    copied_cards = cards.copy()
    l = all_items_on_cards.copy()

    bingo_callouts = []
    bingo_percentage = 0
    while bingo_percentage < percent:
        r = random.choice(l)
        l.remove(r)
        bingo_callouts.append(r)

        delCalloutFromCards(r, copied_cards)
        
        cards_w_bingo = 0
        for card in copied_cards:
            if mode == "line bingo":
                if isLineBingo(card):
                    cards_w_bingo += 1
            if mode == "full card bingo":
                if isFullBingo(card):
                    cards_w_bingo += 1
        bingo_percentage = cards_w_bingo / len(cards) * 100
    return len(bingo_callouts)

def openCalloutFile(callout_file):
    with open(callout_file, "r") as f:
        randomised_item_bank = []
        for line in f.readlines():
            if line[0] != "#" and line != "\n":
                randomised_item_bank.append(line.strip())
    return randomised_item_bank

def calcBingoStats(randomised_item_bank):    
    print("Simulating statistics for current configuration...", end='\r')
    rounds_before_bingo_list = []
    rounds_before_bingo_list_percentage = []
    for _ in range(100):
        all_bingo_cards = []
        for j in range(CARDS_NUMBER):
            bingo_card = generate_bingo_card(randomised_item_bank)
            all_bingo_cards.append(bingo_card)
        rounds = simulateUntilBingo(all_bingo_cards, randomised_item_bank, mode=ANALYSIS_BINGO_MODE)
        rounds_before_bingo_list.append(rounds)

        all_bingo_cards = []
        for j in range(CARDS_NUMBER):
            bingo_card = generate_bingo_card(randomised_item_bank)
            all_bingo_cards.append(bingo_card)
        rounds = simulateUntilBingoPercent(all_bingo_cards, randomised_item_bank, ANALYSIS_PERCENT_UNTIL_BINGO, mode=ANALYSIS_BINGO_MODE)
        rounds_before_bingo_list_percentage.append(rounds)
    
    rounds_before_bingo_list.sort()
    median = statistics.median(rounds_before_bingo_list)
    median_percentage = statistics.median(rounds_before_bingo_list_percentage)
    print(f"⏰ A {ANALYSIS_BINGO_MODE} occurs after ~{int(median)} and at most ~{rounds_before_bingo_list[-1]} callouts. After ~{int(median_percentage)} callouts {ANALYSIS_PERCENT_UNTIL_BINGO}% will have a {ANALYSIS_BINGO_MODE}.")

    # import numpy as np
    # import random
    # from matplotlib import pyplot as plt
    # data = rounds_before_bingo_list

    # bins = np.arange(min(data), max(data), 1) # fixed bin size
    # plt.xlim([0, max(data)+5])
    # plt.hist(data, bins=bins)
    # plt.show()

def showItemBankDlg(item_bank):
    choice = input(f"Press ENTER to show item bank.")
    if choice == "":
        print(f"\nItem bank (sampled {len(item_bank_sample)} from {len(item_bank)}) contains ", end="")
        item_bank_sample.sort()
        for item in item_bank_sample:
            if item == item_bank_sample[-1]:
                print(f"and '{item}'.\n")
            else:
                print(f"'{item}'", end=", ")
    
    choice = input(f"Press ENTER to use this item bank.")
    if choice == "":
        item_bank = item_bank_sample
        return True
    else:
        return False

#Generating cards and randomised callouts
if __name__ == "__main__":
    userConfirmsGeneratedItemBank = False
    while not userConfirmsGeneratedItemBank:

        item_bank = remove_full_duplicates(item_bank)
        report_partial_duplicates(item_bank)

        #Generate item bank
        if NUMBER_OF_ITEMS_TO_CROP_ITEM_BANK_TO > 0:
            if NUMBER_OF_ITEMS_TO_CROP_ITEM_BANK_TO > len(item_bank):
                print(f"Item bank is too small ({len(item_bank)}/{NUMBER_OF_ITEMS_TO_CROP_ITEM_BANK_TO}).")
                exit()
            item_bank_sample = sample(item_bank, NUMBER_OF_ITEMS_TO_CROP_ITEM_BANK_TO)
        else:
            item_bank_sample = item_bank
        for selection in selections_to_prioritise_in_item_bank:
            for item in selection:
                r_index = random.randint(0, len(item_bank_sample) - 1)
                item_bank_sample[r_index] = item

        #Show stats
        calcBingoStats(item_bank_sample)

        #Show item bank
        userConfirmsGeneratedItemBank = showItemBankDlg(item_bank_sample)

    output_file = ""
    while not output_file:
        output_file = filedialog.asksaveasfilename(title="Save bingo", filetypes=[("Bingo files", "*.docx *.txt")])
    
    overwrite = False
    if os.path.exists(output_file):
        overwrite = True
    
    if not overwrite:
        cards_file = f"{output_file}.docx"
        callout_file = f"{output_file}.txt"
    else:
        stem, _ = os.path.splitext(output_file)
        callout_file = f"{stem}.txt"
        cards_file = f"{stem}.docx"

    #Export to Word document
    document = Document()
    paragraph_format = document.styles["Normal"].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    page_margin = Inches(0.3)

    section = document.sections[-1]
    section.page_width = Inches(11.69) # A4
    section.page_height = Inches(8.27)

    all_bingo_cards = []
    for page_num in range(CARDS_NUMBER):
        print(f"Making bingo card #{1 + len(all_bingo_cards)}...")
        # title = document.add_paragraph(f"{card_title}")
        # title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # title_run = title.runs[0]
        # title_run.font.size = Pt(24)
        # title_run.font.name = "Arial"

        table = document.add_table(rows=ROWS, cols=COLUMNS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for column in table.columns:
            for cell in column.cells:
                cell.width = int(floor((section.page_width - (2 * page_margin)) / COLUMNS))
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = int(floor((.875 * section.page_height - (2 * page_margin)) / ROWS)) #TODO: Factor currently hardcoded 

        bingo_card = generate_bingo_card(item_bank)

        # Check if bingo card exists (exact same full card almost impossible)
        while bingo_card in all_bingo_cards:
            print(f"Generated card #{1 + len(all_bingo_cards)} is duplicate of card #{all_bingo_cards.index(bingo_card)}. You got lucky! Retrying...")
            bingo_card = generate_bingo_card(item_bank)
        
        all_bingo_cards.append(bingo_card)

        for i in range(ROWS):
            for j in range(COLUMNS):
                cell = table.cell(i, j)
                MARGIN_INCHES = .2
                set_cell_margins(cell, top=1440*MARGIN_INCHES, start=1440*MARGIN_INCHES, bottom=1440*MARGIN_INCHES, end=1440*MARGIN_INCHES)
                
                if page_num % 2 == 1:
                    if i % 2 == 0 and j % 2 == 0:
                        set_cell_background(cell, COLOR_CELL_BACKGROUND)
                    if i % 2 == 1 and j % 2 == 1:
                        set_cell_background(cell, COLOR_CELL_BACKGROUND)
                else:
                    if i % 2 == 1 and j % 2 == 0:
                        set_cell_background(cell, COLOR_CELL_BACKGROUND)
                    if i % 2 == 0 and j % 2 == 1:
                        set_cell_background(cell, COLOR_CELL_BACKGROUND)

                p = cell.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(str(bingo_card[i][j]))
                run.font.name = "Arial"
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor.from_string(COLOR_FONT)

                run.space_after = Pt(0) # space after paragraph
                set_borders(cell)

        # if page_num < CARDS_NUMBER - 1:
        #     document.add_page_break()

    sections = document.sections
    for section in sections:
        section.top_margin = page_margin
        section.bottom_margin = page_margin
        section.left_margin = page_margin
        section.right_margin = page_margin

    document.save(cards_file)
    
    #Write callout file
    randomised_item_bank = item_bank.copy()
    random.shuffle(randomised_item_bank)
    with open(callout_file, "w", encoding="utf-8") as f:
        for item in randomised_item_bank:
            f.write(f"{item}\n")

    print(f"✔️ Made {CARDS_NUMBER} cards based on {len(item_bank)} items.")